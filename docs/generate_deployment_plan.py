#!/usr/bin/env python3
"""Generate TygaStyle deployment plan as PDF and DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch, mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT_DIR = Path(__file__).resolve().parent
FX = 1420


def usd(n: float) -> str:
    return f"${n:,.0f}"


def rwf(usd_amount: float) -> str:
    return f"{round(usd_amount * FX):,} RWF"


# ---------------------------------------------------------------------------
# Shared content
# ---------------------------------------------------------------------------

TITLE = "TygaStyle — Production Deployment Plan"
SUBTITLE = "Tiger E-commerce hosting proposal · July 2026"
INTRO = (
    "Boss-ready hosting plan for the Tiger E-commerce monorepo: Next.js "
    "storefront + Strapi 5 CMS, Rwanda market (MTN MoMo USSD + WhatsApp). "
    "Prices are vendor list rates as of July 2026; taxes and FX may vary. "
    f"FX assumed ≈ 1,420 RWF per 1 USD."
)

RECOMMENDATION = (
    f"Start with Option B (all on Render, self-hosted Strapi) at about "
    f"{usd(52)}/mo (~{rwf(52)}). It matches the existing Docker/Render path, "
    "stays always-on, and avoids Strapi Cloud sleep on the Starter tier. "
    f"Move to Option A if you want zero CMS ops and managed backups "
    f"(~{usd(97)}/mo)."
)

COMPONENTS = [
    ["Component", "Tech", "Role", "Must stay up?"],
    ["Storefront", "Next.js 16 (Docker)", "Shop, cart, MoMo, WhatsApp", "Yes — public site"],
    ["CMS / Admin", "Strapi 5", "Products, stock, orders, homepage", "Yes — API + staff"],
    ["Database", "PostgreSQL", "Catalog + orders", "Yes"],
    ["Media", "Uploads / CDN", "Product images (~100 MB today)", "Yes"],
    ["Payments", "MTN MoMo *182*…*55066#", "USSD; no card PSP", "N/A (external)"],
    ["Contact", "WhatsApp +250 784 815 151", "Order follow-up links", "N/A (external)"],
]

OPTIONS = [
    ["Option", "Stack", "$ / mo", "≈ RWF / mo", "Year-1 host", "Ops", "Fit"],
    [
        "A · Managed CMS",
        "Render Starter ($7) + Strapi Cloud Pro ($90)",
        usd(97),
        rwf(97),
        usd(97 * 12),
        "Low",
        "Easiest CMS; always-on",
    ],
    [
        "B · Best value ★",
        "Render: Next ($7) + Strapi ($25) + Postgres ($19) + disk (~$1)",
        usd(52),
        rwf(52),
        usd(52 * 12),
        "Medium",
        "Recommended launch",
    ],
    [
        "C · Premium FE",
        "Vercel Pro ($20) + Strapi Cloud Pro ($90)",
        usd(110),
        rwf(110),
        usd(110 * 12),
        "Low",
        "Best CDN / DX for Next",
    ],
]

OPTION_B_LINES = [
    ["Line item", "Spec", "Price"],
    ["Web service — storefront", "Starter · 512 MB", "$7"],
    ["Web service — Strapi", "Standard · 2 GB", "$25"],
    ["Render Postgres", "Basic-1gb", "$19"],
    ["Persistent disk (uploads)", "~5 GB × $0.25", "~$1"],
    ["Workspace", "Hobby", "$0"],
    ["Included bandwidth", "5 GB then $0.15/GB", "usually $0 early"],
]

EXTRAS = [
    ["Item", "Est. cost", "Notes"],
    ["Domain .rw (local registry)", "~15,000 RWF / yr (~$11)", "register.rw / RICTA"],
    ["Domain .com (optional)", "~$12–18 / yr", "Or ~25,000 RWF via register.rw"],
    ["TLS / HTTPS", "$0", "Included on Render / Vercel / Strapi Cloud"],
    ["DNS (Cloudflare free)", "$0", "Optional"],
    ["MTN MoMo merchant", "Existing (code 55066)", "No online gateway fee in current design"],
    ["Email / SMS / Stripe", "$0 now", "Not wired; add later if needed"],
    ["Engineering setup time", "16–24 hours", "Accounts, secrets, deploy, DNS, smoke tests"],
]

YEAR1 = [
    ["Scenario", "Monthly", "Year-1 (host + domain + 15% buffer)"],
    ["Lean (Option B) ★", usd(52), usd(round(52 * 12 * 1.15) + 15)],
    ["Managed (Option A)", usd(97), usd(round(97 * 12 * 1.15) + 15)],
    ["Premium FE (Option C)", usd(110), usd(round(110 * 12 * 1.15) + 15)],
]

TIMELINE = [
    ["Day", "Work", "Owner", "Exit criteria"],
    ["1", "Accounts, billing card, domain, secrets checklist", "Dev + finance", "Render + domain registered"],
    ["2", "Deploy Strapi + Postgres; secrets; seed/import catalog", "Dev", "Admin login; products via API"],
    ["3", "Deploy Next.js; wire CORS FRONTEND_URL", "Dev", "Shop on staging with live catalog"],
    ["4", "DNS + TLS; WhatsApp; MoMo smoke test", "Dev + ops", "HTTPS domain; test order E2E"],
    ["5", "Staff training, backup check, go-live", "Dev + business", "Production live; rollback noted"],
]

TECH_CHECK = [
    "Production secrets rotated (not from backup env files)",
    "NEXT_PUBLIC_STRAPI_URL + SITE_URL set at build time",
    "FRONTEND_URL in Strapi CORS allowlist",
    "DATABASE_SSL on for managed Postgres",
    "Media uploads persist (disk or Cloud CDN)",
    "Admin user created; seed disabled after import",
]

BIZ_CHECK = [
    "Confirm MoMo merchant code 55066",
    "Confirm WhatsApp number on all CTAs",
    "Staff can edit products / mark orders paid",
    "Domain points to storefront only (not admin)",
    "Optional: Cloudflare DNS + basic WAF",
    "Document who pays the monthly invoice",
]

DECISION = (
    f"Approve Option B at ~{usd(52)}/mo (~{rwf(52)}) plus domain, with a "
    f"year-1 ceiling near {usd(round(52 * 12 * 1.15) + 15)} including buffer — "
    f"or choose Option A if managed CMS is worth the extra ~{usd(45)}/mo."
)

FOOTER_NOTE = (
    "Prices exclude VAT/local taxes. Confirm on vendor dashboards before "
    "purchase; Strapi, Render, and Vercel update list prices periodically. "
    "Sources: strapi.io/pricing-cloud, render.com/pricing, vercel.com/pricing, "
    "register.rw (July 2026)."
)

AVOID_NOTE = (
    "Avoid free tiers for real customers (Render Free web, free Postgres, "
    "Strapi Cloud sleep on Starter). Fine for demos; not for a live shop. "
    f"Strapi Cloud Starter is {usd(35)}/mo but sleeps when idle — skip for "
    "production unless you upgrade to Pro."
)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if i == 0 and header:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if i == 0 and header:
                set_cell_shading(cell, "1F4E79")
            elif i % 2 == 0:
                set_cell_shading(cell, "F2F2F2")
    doc.add_paragraph()


def add_heading_styled(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(SUBTITLE)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph(INTRO)

    add_heading_styled(doc, "Executive summary", 1)
    p = doc.add_paragraph()
    r = p.add_run("Recommendation: ")
    r.bold = True
    p.add_run(RECOMMENDATION)

    # KPI-style bullets
    for label, value in [
        ("Monthly hosting (recommended)", f"~{usd(52)}–{usd(110)} depending on option"),
        ("Year-1 hosting (Option B + buffer)", f"~{usd(round(52 * 12 * 1.15) + 15)} (~{rwf(round(52 * 12 * 1.15) + 15)})"),
        ("Go-live timeline", "~5 working days / ~1 week"),
        ("Payment gateway fees (current design)", "$0 — MTN MoMo USSD only"),
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(f"{label}: ").bold = True
        bp.add_run(value)

    add_heading_styled(doc, "1. What we are deploying", 1)
    doc.add_paragraph(
        "Two apps, one database, media storage. No Redis, no Stripe, and no "
        "email provider are required by the current codebase."
    )
    add_table(doc, COMPONENTS)

    add_heading_styled(doc, "2. Target architecture", 1)
    doc.add_paragraph(
        "Browser → Next.js storefront (HTTPS) → Strapi REST API → PostgreSQL + media storage."
    )
    for title_t, body in [
        ("Storefront (public)", "Next.js standalone on port 3000. Custom domain + TLS. Build-time env for Strapi URL and site URL."),
        ("CMS (staff)", "Strapi 5 admin + REST API. CORS locked to shop domain. Secrets for JWT / APP_KEYS."),
        ("Data", "Managed Postgres + disk or CDN. Start small (1 GB RAM DB, 5–50 GB media); scale later."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{title_t}: ").bold = True
        p.add_run(body)

    add_heading_styled(doc, "3. Hosting options compared", 1)
    doc.add_paragraph(
        "Monthly run-rate for a small boutique (low–medium traffic). Domain and "
        "one-time setup are listed separately."
    )
    add_table(doc, OPTIONS)

    warn = doc.add_paragraph()
    r = warn.add_run("Important: ")
    r.bold = True
    warn.add_run(AVOID_NOTE)

    add_heading_styled(doc, "4. Option B detail (recommended)", 1)
    doc.add_paragraph(
        "Pros: always-on, predictable bill, uses existing Dockerfile. "
        "Cons: you own backups, upgrades, and disk growth. Paid Render Postgres "
        "includes a point-in-time recovery window."
    )
    add_table(doc, OPTION_B_LINES)

    add_heading_styled(doc, "5. Options A and C (summary)", 1)
    p = doc.add_paragraph()
    p.add_run("Option A — Managed CMS: ").bold = True
    p.add_run(
        "Render Starter for Next.js + Strapi Cloud Pro (1M API requests, 250 GB "
        "assets, weekly backups, always-on). Best if nobody wants to babysit "
        "Postgres. Higher fixed cost."
    )
    p = doc.add_paragraph()
    p.add_run("Option C — Premium frontend: ").bold = True
    p.add_run(
        "Vercel Pro seat ($20 + usage credit) for Next.js + Strapi Cloud Pro. "
        "Best global CDN and preview deploys; slightly more expensive until "
        "traffic or marketing needs it."
    )

    add_heading_styled(doc, "6. One-time & annual extras", 1)
    add_table(doc, EXTRAS)

    add_heading_styled(doc, "7. Year-1 budget", 1)
    doc.add_paragraph(
        f"Recommended Option B: hosting {usd(52 * 12)}/year + domain ~$15 + 15% "
        f"contingency ≈ {usd(round(52 * 12 * 1.15) + 15)} "
        f"(~{rwf(round(52 * 12 * 1.15) + 15)})."
    )
    add_table(doc, YEAR1)

    add_heading_styled(doc, "8. Launch timeline (~5 working days)", 1)
    add_table(doc, TIMELINE)

    add_heading_styled(doc, "9. Go-live checklist", 1)
    add_heading_styled(doc, "Technical", 2)
    for item in TECH_CHECK:
        doc.add_paragraph(item, style="List Bullet")
    add_heading_styled(doc, "Business", 2)
    for item in BIZ_CHECK:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_styled(doc, "10. Decision ask", 1)
    doc.add_paragraph(DECISION)
    note = doc.add_paragraph()
    r = note.add_run(FOOTER_NOTE)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)

    sign = doc.add_paragraph()
    sign.add_run("\nPrepared for internal review · Tiger E-commerce / TygaStyle").italic = True

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

NAVY = colors.HexColor("#1F4E79")
LIGHT = colors.HexColor("#F2F2F2")
MUTED = colors.HexColor("#595959")
ACCENT = colors.HexColor("#2E7D32")


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="DocTitle",
            parent=styles["Title"],
            fontSize=18,
            textColor=NAVY,
            spaceAfter=4,
            alignment=TA_CENTER,
            leading=22,
        )
    )
    styles.add(
        ParagraphStyle(
            name="DocSub",
            parent=styles["Normal"],
            fontSize=10,
            textColor=MUTED,
            alignment=TA_CENTER,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Section",
            parent=styles["Heading1"],
            fontSize=13,
            textColor=NAVY,
            spaceBefore=14,
            spaceAfter=6,
            leading=16,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SubSection",
            parent=styles["Heading2"],
            fontSize=11,
            textColor=NAVY,
            spaceBefore=8,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Cell",
            parent=styles["Normal"],
            fontSize=8,
            leading=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CellHead",
            parent=styles["Normal"],
            fontSize=8,
            leading=10,
            textColor=colors.white,
            fontName="Helvetica-Bold",
        )
    )
    styles.add(
        ParagraphStyle(
            name="FooterNote",
            parent=styles["Normal"],
            fontSize=8,
            textColor=MUTED,
            leading=11,
            spaceBefore=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Callout",
            parent=styles["Normal"],
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#0D47A1"),
            leftIndent=6,
            rightIndent=6,
            spaceBefore=4,
            spaceAfter=8,
        )
    )
    return styles


def pdf_table(data: list[list[str]], styles, col_widths=None) -> Table:
    styled = []
    for i, row in enumerate(data):
        style = styles["CellHead"] if i == 0 else styles["Cell"]
        styled.append([Paragraph(str(c).replace("\n", "<br/>"), style) for c in row])
    t = Table(styled, colWidths=col_widths, repeatRows=1)
    cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            cmds.append(("BACKGROUND", (0, i), (-1, i), LIGHT))
    t.setStyle(TableStyle(cmds))
    return t


def build_pdf(path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=TITLE,
        author="Tiger E-commerce / TygaStyle",
    )
    story = []
    story.append(Paragraph(TITLE, styles["DocTitle"]))
    story.append(Paragraph(SUBTITLE, styles["DocSub"]))
    story.append(Paragraph(INTRO, styles["Body"]))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=8))

    story.append(Paragraph("Executive summary", styles["Section"]))
    story.append(Paragraph(f"<b>Recommendation:</b> {RECOMMENDATION}", styles["Callout"]))
    for label, value in [
        ("Monthly hosting (recommended)", f"~{usd(52)}–{usd(110)} depending on option"),
        ("Year-1 hosting (Option B + buffer)", f"~{usd(round(52 * 12 * 1.15) + 15)} (~{rwf(round(52 * 12 * 1.15) + 15)})"),
        ("Go-live timeline", "~5 working days / ~1 week"),
        ("Payment gateway fees (current design)", "$0 — MTN MoMo USSD only"),
    ]:
        story.append(Paragraph(f"• <b>{label}:</b> {value}", styles["Body"]))

    story.append(Paragraph("1. What we are deploying", styles["Section"]))
    story.append(
        Paragraph(
            "Two apps, one database, media storage. No Redis, no Stripe, and no "
            "email provider are required by the current codebase.",
            styles["Body"],
        )
    )
    story.append(pdf_table(COMPONENTS, styles, col_widths=[72, 95, 160, 95]))

    story.append(Paragraph("2. Target architecture", styles["Section"]))
    story.append(
        Paragraph(
            "Browser → Next.js storefront (HTTPS) → Strapi REST API → PostgreSQL + media storage.",
            styles["Body"],
        )
    )
    for title_t, body in [
        ("Storefront (public)", "Next.js standalone on port 3000. Custom domain + TLS. Build-time env for Strapi URL and site URL."),
        ("CMS (staff)", "Strapi 5 admin + REST API. CORS locked to shop domain. Secrets for JWT / APP_KEYS."),
        ("Data", "Managed Postgres + disk or CDN. Start small (1 GB RAM DB, 5–50 GB media); scale later."),
    ]:
        story.append(Paragraph(f"• <b>{title_t}:</b> {body}", styles["Body"]))

    story.append(Paragraph("3. Hosting options compared", styles["Section"]))
    story.append(
        Paragraph(
            "Monthly run-rate for a small boutique (low–medium traffic). Domain and "
            "one-time setup are listed separately.",
            styles["Body"],
        )
    )
    story.append(pdf_table(OPTIONS, styles, col_widths=[68, 130, 42, 58, 52, 36, 70]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>Important:</b> {AVOID_NOTE}", styles["Body"]))

    story.append(Paragraph("4. Option B detail (recommended)", styles["Section"]))
    story.append(
        Paragraph(
            "Pros: always-on, predictable bill, uses existing Dockerfile. "
            "Cons: you own backups, upgrades, and disk growth. Paid Render Postgres "
            "includes a point-in-time recovery window.",
            styles["Body"],
        )
    )
    story.append(pdf_table(OPTION_B_LINES, styles, col_widths=[180, 160, 80]))

    story.append(Paragraph("5. Options A and C (summary)", styles["Section"]))
    story.append(
        Paragraph(
            "<b>Option A — Managed CMS:</b> Render Starter for Next.js + Strapi Cloud Pro "
            "(1M API requests, 250 GB assets, weekly backups, always-on). Best if nobody "
            "wants to babysit Postgres. Higher fixed cost.",
            styles["Body"],
        )
    )
    story.append(
        Paragraph(
            "<b>Option C — Premium frontend:</b> Vercel Pro seat ($20 + usage credit) for "
            "Next.js + Strapi Cloud Pro. Best global CDN and preview deploys; slightly more "
            "expensive until traffic or marketing needs it.",
            styles["Body"],
        )
    )

    story.append(Paragraph("6. One-time & annual extras", styles["Section"]))
    story.append(pdf_table(EXTRAS, styles, col_widths=[150, 130, 160]))

    story.append(Paragraph("7. Year-1 budget", styles["Section"]))
    story.append(
        Paragraph(
            f"Recommended Option B: hosting {usd(52 * 12)}/year + domain ~$15 + 15% "
            f"contingency ≈ <b>{usd(round(52 * 12 * 1.15) + 15)}</b> "
            f"(~{rwf(round(52 * 12 * 1.15) + 15)}).",
            styles["Body"],
        )
    )
    story.append(pdf_table(YEAR1, styles, col_widths=[160, 80, 200]))

    story.append(Paragraph("8. Launch timeline (~5 working days)", styles["Section"]))
    story.append(pdf_table(TIMELINE, styles, col_widths=[30, 175, 70, 145]))

    story.append(Paragraph("9. Go-live checklist", styles["Section"]))
    story.append(Paragraph("Technical", styles["SubSection"]))
    for item in TECH_CHECK:
        story.append(Paragraph(f"• {item}", styles["Body"]))
    story.append(Paragraph("Business", styles["SubSection"]))
    for item in BIZ_CHECK:
        story.append(Paragraph(f"• {item}", styles["Body"]))

    story.append(Paragraph("10. Decision ask", styles["Section"]))
    story.append(Paragraph(DECISION, styles["Body"]))
    story.append(Paragraph(FOOTER_NOTE, styles["FooterNote"]))
    story.append(
        Paragraph(
            "<i>Prepared for internal review · Tiger E-commerce / TygaStyle</i>",
            styles["FooterNote"],
        )
    )

    doc.build(story)


def main() -> None:
    pdf_path = OUT_DIR / "TygaStyle-Deployment-Plan.pdf"
    docx_path = OUT_DIR / "TygaStyle-Deployment-Plan.docx"
    build_pdf(pdf_path)
    build_docx(docx_path)
    print(f"Wrote {pdf_path}")
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
